"""
Text Extractor Service - Extract text from various document formats
Uses LlamaParse as the primary document loader for superior parsing,
with fallback to local extractors (PyPDF, python-docx).
"""

import io
import os
import tempfile
import asyncio
from typing import Optional
from fastapi import HTTPException


class TextExtractorService:
    """
    Service to extract text from different document formats.

    Primary: LlamaParse (cloud-based, handles complex PDFs with tables/images)
    Fallback: Local extractors (PyPDF for PDF, python-docx for DOCX, raw decode for TXT/MD)

    Supports: PDF, DOCX, DOC, TXT, MD, PPTX, XLSX, HTML, and more via LlamaParse
    """

    # Additional formats supported by LlamaParse beyond our original set
    LLAMAPARSE_SUPPORTED = {
        ".pdf",
        ".docx",
        ".doc",
        ".pptx",
        ".xlsx",
        ".html",
        ".htm",
        ".txt",
        ".md",
        ".csv",
        ".epub",
        ".rtf",
    }

    def __init__(self):
        self._llama_api_key = os.getenv("LLAMA_CLOUD_API_KEY")
        self._use_llamaparse = bool(self._llama_api_key)

        if self._use_llamaparse:
            print("[TextExtractor] LlamaParse enabled (LLAMA_CLOUD_API_KEY found)")
        else:
            print("[TextExtractor] LlamaParse disabled - using local extractors")
            print(
                "[TextExtractor] Set LLAMA_CLOUD_API_KEY in .env to enable LlamaParse"
            )

    async def extract(self, content: bytes, filename: str) -> str:
        """
        Extract text from document based on file extension.
        Uses LlamaParse if available, otherwise falls back to local extractors.

        Args:
            content: Raw bytes of the file
            filename: Name of the file (used to determine type)

        Returns:
            Extracted text as string
        """
        filename = filename.lower()

        # Try LlamaParse first if configured
        if self._use_llamaparse:
            try:
                text = await self._extract_with_llamaparse(content, filename)
                if text and text.strip():
                    print(
                        f"[TextExtractor] Successfully parsed '{filename}' with LlamaParse"
                    )
                    return text
                else:
                    print(
                        f"[TextExtractor] LlamaParse returned empty text for '{filename}', trying fallback"
                    )
            except Exception as e:
                print(f"[TextExtractor] LlamaParse failed for '{filename}': {e}")
                print("[TextExtractor] Falling back to local extractors...")

        # Fallback to local extractors
        return await self._extract_local(content, filename)

    # ============ LlamaParse Extraction ============

    async def _extract_with_llamaparse(self, content: bytes, filename: str) -> str:
        """
        Extract text using LlamaParse cloud API.

        LlamaParse excels at:
        - Complex PDF layouts with tables, images, and multi-column text
        - Preserving document structure and formatting
        - Extracting text from scanned/image-based PDFs (OCR)
        - Handling diverse document formats
        """
        try:
            from llama_parse import LlamaParse
        except ImportError:
            raise HTTPException(
                status_code=500,
                detail="llama-parse package not installed. Run: pip install llama-parse",
            )

        # LlamaParse requires a file path, so we write to a temp file
        file_ext = self._get_extension(filename)

        # Create temp file with the correct extension
        tmp_file = None
        try:
            tmp_file = tempfile.NamedTemporaryFile(
                suffix=file_ext, delete=False, prefix="llamaparse_"
            )
            tmp_file.write(content)
            tmp_file.flush()
            tmp_file.close()

            # Configure LlamaParse with optimal settings
            parser = LlamaParse(
                api_key=self._llama_api_key,
                result_type="markdown",  # Get markdown-formatted output
                verbose=False,
                language="en",
                # Parsing quality settings
                num_workers=1,  # Single worker for individual files
            )

            # Parse the document - run in executor since LlamaParse may block
            documents = await asyncio.get_event_loop().run_in_executor(
                None, lambda: parser.load_data(tmp_file.name)
            )

            if not documents:
                return ""

            # Combine all document pages/sections into one text
            text_parts = []
            for i, doc in enumerate(documents):
                page_text = doc.text if hasattr(doc, "text") else str(doc)
                if page_text and page_text.strip():
                    text_parts.append(page_text)

            return "\n\n".join(text_parts)

        finally:
            # Clean up temp file
            if tmp_file and os.path.exists(tmp_file.name):
                try:
                    os.unlink(tmp_file.name)
                except OSError:
                    pass

    # ============ Local Fallback Extractors ============

    async def _extract_local(self, content: bytes, filename: str) -> str:
        """Route to appropriate local extractor based on file extension."""
        if filename.endswith(".pdf"):
            return await self._extract_pdf(content)
        elif filename.endswith(".docx"):
            return await self._extract_docx(content)
        elif filename.endswith(".doc"):
            return await self._extract_doc(content)
        elif filename.endswith((".txt", ".md")):
            return await self._extract_text(content)
        else:
            raise HTTPException(
                status_code=400, detail=f"Unsupported file format: {filename}"
            )

    async def _extract_pdf(self, content: bytes) -> str:
        """Extract text from PDF using PyPDF (fallback)"""
        try:
            from pypdf import PdfReader

            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)

            text_parts = []
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")

            return "\n\n".join(text_parts)

        except ImportError:
            raise HTTPException(
                status_code=500, detail="PDF processing library (pypdf) not installed"
            )
        except Exception as e:
            raise HTTPException(
                status_code=400, detail=f"Failed to extract text from PDF: {str(e)}"
            )

    async def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX using python-docx (fallback)"""
        try:
            from docx import Document

            docx_file = io.BytesIO(content)
            doc = Document(docx_file)

            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)

        except ImportError:
            raise HTTPException(
                status_code=500,
                detail="DOCX processing library (python-docx) not installed",
            )
        except Exception as e:
            raise HTTPException(
                status_code=400, detail=f"Failed to extract text from DOCX: {str(e)}"
            )

    async def _extract_doc(self, content: bytes) -> str:
        """Extract text from legacy DOC format"""
        raise HTTPException(
            status_code=400,
            detail="Legacy .doc format not fully supported. Please convert to .docx or .pdf",
        )

    async def _extract_text(self, content: bytes) -> str:
        """Extract text from plain text files (TXT, MD)"""
        try:
            encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]

            for encoding in encodings:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue

            # If all encodings fail, use UTF-8 with error replacement
            return content.decode("utf-8", errors="replace")

        except Exception as e:
            raise HTTPException(
                status_code=400, detail=f"Failed to read text file: {str(e)}"
            )

    # ============ Utility ============

    @staticmethod
    def _get_extension(filename: str) -> str:
        """Get the file extension from filename."""
        if "." in filename:
            return "." + filename.rsplit(".", 1)[-1].lower()
        return ""
